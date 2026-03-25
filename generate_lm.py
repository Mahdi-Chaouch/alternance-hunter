#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Génération de lettres de motivation (LM) à partir du fichier draft et d'un template .docx.

Placeholders supportés dans le template :
  {{ENTREPRISE}}  - Nom de l'entreprise
  {{DATE}}        - Date du jour (jj/mm/aaaa)
  {{ZONE}}        - Zone géographique (ex: Paris (Paris-Centre))
  {{VILLE}}       - Ville (ex: Paris, Cannes, Fontainebleau, Auxerre)
  {{DOMAINE}}     - Domaine du site (ex: exemple.com)
  {{SITE}}        - URL du site
  {{SECTEUR}}     - Secteur d'activité (ex: Informatique / Digital)
  {{SPECIALITE}}  - Métier / spécialité (ex: Développement web)
  {{PARAGRAPHE_PERSONNALISE}} - Paragraphe personnalisé (vide si l'IA n'est pas activée)
"""

import argparse
import re
import sys
from pathlib import Path
from datetime import date
from typing import Tuple
from urllib.parse import urlparse

from docx import Document

SEP = "=============================="

# Libellés secteur pour normaliser l'affichage quand la spécialité est vide.
SECTOR_LABELS_FOR_PROMPT: dict[str, str] = {
    "it": "Informatique / Digital",
    "food": "Alimentation / Restauration",
    "law": "Droit / Finance / Assurance",
    "trade": "Commerce / Retail",
    "health": "Santé / Médical",
    "construction": "BTP / Construction / Artisanat",
    "marketing": "Commerce / Marketing",
    "finance": "Finance",
    "all": "tous secteurs",
}


def safe_filename(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]', "", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name[:120]


def extract_domain(site: str) -> str:
    """Extrait le domaine (sans www) depuis une URL."""
    if not site:
        return ""
    try:
        netloc = urlparse(site.strip()).netloc or ""
        return netloc.replace("www.", "").lower().strip()
    except Exception:
        return ""


def extract_company_info(draft_file: str) -> list[dict]:
    """
    Parse le fichier draft et retourne une liste de dicts uniques par entreprise :
    { "company", "zone", "site", "domain" }
    """
    text = Path(draft_file).read_text(encoding="utf-8", errors="ignore")
    blocks = [b.strip() for b in text.split(SEP) if b.strip()]

    seen = set()
    result = []

    for block in blocks:
        m_company = re.search(r"Entreprise:\s*(.+)", block, re.IGNORECASE)
        m_zone = re.search(r"Zone:\s*(.+)", block)
        m_ville = re.search(r"Ville:\s*(.+)", block)
        m_site = re.search(r"Site:\s*(.+)", block)

        if not m_company:
            continue

        company = m_company.group(1).strip()
        key = company.lower()
        if key in seen:
            continue
        seen.add(key)

        zone = m_zone.group(1).strip() if m_zone else ""
        ville = m_ville.group(1).strip() if m_ville else ""
        site = m_site.group(1).strip() if m_site else ""
        domain = extract_domain(site)

        result.append({
            "company": company,
            "zone": zone,
            "ville": ville,
            "site": site,
            "domain": domain,
        })

    return result


def _replace_in_paragraph(p, replacements: dict) -> None:
    """Remplace les placeholders dans un paragraphe (conserve le premier run pour le texte)."""
    text = p.text
    if not any(ph in text for ph in replacements):
        return
    for placeholder, value in replacements.items():
        text = text.replace(placeholder, str(value or ""))
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text


def _replace_in_cell(cell, replacements: dict) -> None:
    for p in cell.paragraphs:
        _replace_in_paragraph(p, replacements)


def replace_in_doc(
    doc: Document,
    info: dict,
    sector: str = "",
    specialty: str = "",
) -> None:
    """
    Remplace tous les placeholders dans le document.
    info : { company, zone, site, domain, ville }
    sector / specialty : contexte run pour {{SECTEUR}} et {{SPECIALITE}}
    """
    today = date.today().strftime("%d/%m/%Y")
    replacements = {
        "{{ENTREPRISE}}": info.get("company", ""),
        "{{DATE}}": today,
        "{{ZONE}}": info.get("zone", ""),
        "{{VILLE}}": info.get("ville", ""),
        "{{DOMAINE}}": info.get("domain", ""),
        "{{SITE}}": info.get("site", ""),
        "{{SECTEUR}}": sector,
        "{{SPECIALITE}}": specialty,
    }
    # L'IA est retirée du produit : on remplace toujours ce placeholder par vide
    # pour éviter que le token brut reste dans le document.
    replacements["{{PARAGRAPHE_PERSONNALISE}}"] = ""

    for p in doc.paragraphs:
        _replace_in_paragraph(p, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)


def main() -> None:
    parser = argparse.ArgumentParser(description="Génère des LM .docx à partir du draft et du template.")
    parser.add_argument("--draft-file", default="data/exports/draft_emails.txt", help="Fichier draft (blocs Entreprise/Zone/Site)")
    parser.add_argument("--template", default="", help="Template Word avec placeholders (obligatoire; uploadez via l'UI ou --template /chemin)")
    parser.add_argument("--out-dir", default="outputs/letters", help="Dossier de sortie des LM")
    parser.add_argument("--sector", default="", help="Secteur d'activité pour placeholders")
    parser.add_argument("--specialty", default="", help="Métier / spécialité pour placeholders")

    args = parser.parse_args()

    draft = Path(args.draft_file)
    template = Path(args.template)
    outdir = Path(args.out_dir)

    if not draft.exists():
        raise SystemExit(f"❌ Fichier draft introuvable: {draft} (cwd={Path.cwd()})")

    if not args.template.strip() or not template.exists():
        raise SystemExit(
            "❌ Template LM requis. Uploadez un template (.docx) depuis l'interface web "
            "ou fournissez --template /chemin/vers/template.docx"
        )

    outdir.mkdir(parents=True, exist_ok=True)

    companies_info = extract_company_info(str(draft))
    print("Entreprises trouvées:", len(companies_info))
    if not companies_info:
        raise SystemExit(
            "❌ Aucun bloc entreprise valide trouvé dans le fichier draft.\n"
            "   Lancez d'abord l'étape « Recherche d'entreprises » (hunter) pour générer des brouillons "
            "avec des entreprises ayant un email trouvé, puis relancez la génération des lettres."
        )

    sector = getattr(args, "sector", "") or ""
    specialty = getattr(args, "specialty", "") or ""
    sector_display = SECTOR_LABELS_FOR_PROMPT.get(sector, sector) or sector

    generated = 0
    for info in companies_info:
        doc = Document(template)
        replace_in_doc(doc, info, sector=sector_display, specialty=specialty)

        filename = safe_filename(info["company"]) + "_LM.docx"
        path = outdir / filename
        doc.save(path)
        generated += 1

    print("LM générées :", generated)
    print("Dossier :", outdir)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n⛔ Interrompu.")
        sys.exit(1)
